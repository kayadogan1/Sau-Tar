from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "rapor" / "tarsau_proje_raporu.docx"

BLUE = "2E74B5"
DARK_BLUE = "0B2545"
MID_BLUE = "1F4D78"
MUTED = "555555"
BORDER = "D9E2EC"
LIGHT_FILL = "F2F4F7"
CODE_FILL = "F4F6F9"
PASS_FILL = "EAF7EF"


def set_run_font(run, font="Calibri", size=11, bold=False, color=None):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, inches):
    width = int(inches * 1440)
    cell.width = Inches(inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(table, top=90, start=130, bottom=90, end=130):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def table_fixed_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_para(doc, text="", size=11, bold=False, color=None, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    set_run_font(
        run,
        size=16 if level == 1 else 13 if level == 2 else 12,
        bold=True,
        color=BLUE if level < 3 else MID_BLUE,
    )
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run)


def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table_fixed_width(table, [6.3])
    set_table_borders(table)
    set_cell_margins(table, top=130, start=160, bottom=130, end=160)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CODE_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for index, line in enumerate(text.rstrip().splitlines()):
        if index:
            p.add_run("\n")
        run = p.add_run(line)
        set_run_font(run, font="Courier New", size=8.5)
    add_para(doc, after=2)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table_fixed_width(table, [1.7, 4.6])
    set_table_borders(table)
    set_cell_margins(table)
    for idx, (key, value) in enumerate(rows):
        left, right = table.rows[idx].cells
        set_cell_shading(left, LIGHT_FILL)
        p1 = left.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        set_run_font(p1.add_run(key), bold=True)
        p2 = right.paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        set_run_font(p2.add_run(value))
    add_para(doc, after=2)


def add_matrix_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table_fixed_width(table, widths)
    set_table_borders(table)
    set_cell_margins(table, top=75, start=95, bottom=75, end=95)

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(header), size=9.5, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            if str(value).upper() == "PASS":
                set_cell_shading(cell, PASS_FILL)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if i in (0, len(row) - 1):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(p.add_run(str(value)), size=8.8, bold=(str(value).upper() == "PASS"))
    add_para(doc, after=2)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, MID_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("Sistem Programlama Proje Raporu - tarsau")
    set_run_font(run, size=9, color="666666")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("tarsau .sau arşivleme programı")
    set_run_font(run, size=9, color="777777")


def add_cover(doc):
    add_para(doc, "BİLGİSAYAR MÜHENDİSLİĞİ", size=10.5, bold=True, color=MUTED, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    title = add_para(doc, "tarsau Arşivleme Programı", size=24, bold=True, color=DARK_BLUE, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    title.paragraph_format.space_before = Pt(18)
    add_para(
        doc,
        "Sistem Programlama 2025-2026 Bahar Dönemi Proje Raporu",
        size=12.5,
        color=MUTED,
        after=18,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_kv_table(
        doc,
        [
            ("Proje", "Sıkıştırma yapmadan ASCII metin dosyalarını .sau arşivinde birleştirme ve geri açma"),
            ("Program", "tarsau"),
            ("Dil ve ortam", "C, Linux/Unix uyumlu komut satırı, Makefile"),
            ("Kaynak depo", "Bu çalışma alanında .git/GitHub remote bilgisi bulunmadı; sahte link eklenmedi"),
            ("Teslim tarihi", "24 Mayıs 2026 23:59"),
            ("Test özeti", "22 uçtan uca senaryo ve make test başarıyla tamamlandı"),
            ("Rapor tarihi", "19 Mayıs 2026"),
        ],
    )


def build():
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_heading(doc, "1. Amaç ve Kapsam")
    add_para(
        doc,
        "Bu çalışmada tar/rar/zip benzeri çalışan fakat sıkıştırma yapmayan tarsau adlı komut satırı programı geliştirilmiştir. "
        "Program -b modu ile metin dosyalarını tek bir .sau arşivinde toplar, -a modu ile arşivi istenen dizine açar. "
        "Uygulama, ödev föyünde verilen ASCII dosya koşulu, 32 dosya sınırı, 200 MB toplam boyut sınırı, bozuk arşiv kontrolü "
        "ve dosya izinlerinin korunması gereksinimlerini kapsar."
    )

    add_heading(doc, "2. Gereksinim Uyumu")
    add_matrix_table(
        doc,
        ["No", "Gereksinim", "Uygulamadaki karşılık", "Durum"],
        [
            ("1", "Komut satırında -b ile arşiv oluşturma", "./tarsau -b dosya1 dosya2 ... -o arsiv.sau biçimi desteklenir.", "PASS"),
            ("2", "-o verilmezse a.sau üretme", "Çıktı yolu varsayılan olarak a.sau atanır.", "PASS"),
            ("3", "Yalnızca ASCII metin dosyalarını kabul etme", "NUL ve 127 üzeri baytlar uyumsuz format sayılır.", "PASS"),
            ("4", "En fazla 32 giriş dosyası", "MAX_FILES sabiti 32'dir; fazlası hata üretir.", "PASS"),
            ("5", "Toplam giriş boyutu en fazla 200 MB", "MAX_TOTAL_SIZE sınırı her dosya sonrası kontrol edilir.", "PASS"),
            ("6", ".sau organizasyon bölümünü üretme", "İlk 10 bayt organizasyon uzunluğu, devamında |ad,izin,boyut| kayıtları yazılır.", "PASS"),
            ("7", "-a ile arşiv açma", "./tarsau -a arsiv.sau [dizin] biçimi desteklenir.", "PASS"),
            ("8", "Dizin yoksa oluşturma", "mkdir_p göreceli, mutlak ve iç içe dizinleri oluşturur.", "PASS"),
            ("9", "Dosya izinlerini koruma", "stat ile izin okunur, açma sonunda chmod ile geri yüklenir.", "PASS"),
            ("10", "Bozuk veya uygunsuz arşivde kontrollü çıkış", "Uzantı, başlık, metadata ve toplam boyut doğrulanır.", "PASS"),
        ],
        [0.45, 1.95, 3.2, 0.7],
    )

    add_heading(doc, "3. Kullanım")
    add_heading(doc, "3.1 Arşiv Oluşturma", level=2)
    add_code_block(doc, "$ ./tarsau -b t1 t2 t3 t4.txt t5.dat -o s1.sau\nDosyalar birleştirildi.")
    add_para(doc, "-o parametresi verilmediğinde program arşivi varsayılan olarak a.sau adıyla oluşturur.")
    add_heading(doc, "3.2 Arşiv Açma", level=2)
    add_code_block(doc, "$ ./tarsau -a s1.sau d1\nd1 dizininde t1, t2, t3, t4.txt ve t5.dat dosyaları açıldı.")
    add_para(doc, "Dizin parametresi verilmediğinde dosyalar geçerli dizine açılır. Dizin adı göreceli veya mutlak olabilir.")

    add_heading(doc, "4. Arşiv Formatı")
    add_para(
        doc,
        ".sau dosyası iki bölümden oluşur. İlk bölüm organizasyon bilgisidir; ilk 10 bayt bu bölümün toplam uzunluğunu ASCII sayı olarak tutar. "
        "Bu uzunluğa 10 baytlık başlık alanı da dahildir. Devamında her dosya için |dosya_adı,izin,boyut| kaydı bulunur. "
        "İkinci bölümde dosya içerikleri ayırıcı karakter kullanılmadan, metadata kayıtlarındaki boyut değerleriyle çözülecek şekilde art arda yer alır."
    )
    add_code_block(doc, "0000000064|t1,0644,52||t2,0644,70|<t1 verisi><t2 verisi>")

    add_heading(doc, "5. Tasarım ve Kod Yapısı")
    add_bullet(doc, "InputFile yapısı arşivlenecek dosyanın yolu, güvenli dosya adı, izinleri ve boyutunu taşır.")
    add_bullet(doc, "ArchiveEntry yapısı arşiv açılırken metadata kayıtlarından okunan dosya adı, izin ve boyut bilgilerini taşır.")
    add_bullet(doc, "copy_stream fonksiyonu dosya içeriklerini 64 KB tamponla kopyalar; böylece büyük dosyalar belleğe tamamen alınmaz.")
    add_bullet(doc, "is_safe_archive_name fonksiyonu dosya adlarında ayırıcıları ve dizin kaçışı oluşturabilecek örüntüleri engeller.")
    add_bullet(doc, "parse_metadata fonksiyonu organizasyon bölümünü kayıt kayıt ayrıştırır ve kayıt biçimi bozuksa işlemi durdurur.")

    add_heading(doc, "6. Önemli Kod Parçacıkları")
    add_heading(doc, "6.1 ASCII ve Boyut Denetimi", level=2)
    add_code_block(
        doc,
        """static int validate_ascii_text_file(const char *path, long long *size_out) {
    ...
    if (buffer[i] == 0 || buffer[i] > 127) {
        fclose(file);
        return 1;
    }
    size += (long long) got;
    ...
}""",
    )
    add_heading(doc, "6.2 Metadata Kaydı ve Başlık Üretimi", level=2)
    add_code_block(
        doc,
        """snprintf(record, sizeof(record), "|%s,%04o,%lld|",
         file->name, (unsigned int) (file->mode & 0777), file->size);

snprintf(header, sizeof(header), "%010zu", HEADER_WIDTH + metadata_length);""",
    )
    add_heading(doc, "6.3 Arşiv Açma ve İzinleri Geri Yükleme", level=2)
    add_code_block(
        doc,
        """if (copy_stream(archive, out, entries[i].size) != 0) {
    fprintf(stderr, "Arsiv acilirken hata olustu!\\n");
    ...
}
fclose(out);
chmod(output_path, entries[i].mode);""",
    )

    add_heading(doc, "7. Test Süreci")
    add_para(
        doc,
        "Testler önce Makefile üzerindeki standart hedefle, ardından geçici bir çalışma dizininde uçtan uca senaryo matrisiyle yürütülmüştür. "
        "Senaryolar başarılı arşivleme/açma akışlarını, varsayılan parametreleri, izin korumayı ve beklenen hata durumlarını kapsar."
    )
    add_heading(doc, "7.1 make test Çıktısı", level=2)
    add_code_block(
        doc,
        """$ make test
rm -rf testler/cikti testler/ornek.sau
./tarsau -b testler/t1 testler/t2 testler/t3 testler/t4.txt testler/t5.dat -o testler/ornek.sau
Dosyalar birleştirildi.
./tarsau -a testler/ornek.sau testler/cikti
testler/cikti dizininde t1, t2, t3, t4.txt ve t5.dat dosyaları açıldı.
diff -q testler/t1 testler/cikti/t1
diff -q testler/t2 testler/cikti/t2
diff -q testler/t3 testler/cikti/t3
diff -q testler/t4.txt testler/cikti/t4.txt
diff -q testler/t5.dat testler/cikti/t5.dat""",
    )
    add_heading(doc, "7.2 Uçtan Uca Test Matrisi", level=2)
    add_matrix_table(
        doc,
        ["No", "Test senaryosu", "Beklenen sonuç", "Durum"],
        [
            ("1", "5 dosya ile arşiv oluşturma ve ayrı dizine açma", "Tüm dosyalar diff ile birebir aynı.", "PASS"),
            ("2", "-o verilmeden arşiv oluşturma", "a.sau dosyası oluşur.", "PASS"),
            ("3", "Dizin verilmeden arşiv açma", "Dosya geçerli dizine açılır.", "PASS"),
            ("4", "Göreceli iç içe dizine açma", "Hedef dizinler oluşturulur, içerik korunur.", "PASS"),
            ("5", "Mutlak dizine açma", "Mutlak hedefte dosya oluşur ve içerik aynıdır.", "PASS"),
            ("6", "Çalıştırma izni olan dosyayı arşivleme", "Açılan dosyanın izni kaynakla aynıdır.", "PASS"),
            ("7", "Türkçe karakter içeren giriş dosyası", "Uyumsuz format mesajı ile reddedilir.", "PASS"),
            ("8", "NUL baytı içeren giriş dosyası", "Uyumsuz format mesajı ile reddedilir.", "PASS"),
            ("9", "-b sonrası giriş dosyası vermeme", "Kullanım mesajı ile kontrollü çıkış.", "PASS"),
            ("10", "-o sonrası arşiv adı vermeme", "Kullanım mesajı ile kontrollü çıkış.", "PASS"),
            ("11", "Olmayan giriş dosyası", "Okunamadı hatası ile çıkış.", "PASS"),
            ("12", "Aynı basename'e sahip iki dosya", "Arşiv içi ad tekrarı reddedilir.", "PASS"),
            ("13", "Virgül içeren dosya adı", "Arşiv formatı için uygunsuz ad reddedilir.", "PASS"),
            ("14", "33 giriş dosyası", "En fazla 32 dosya sınırı uygulanır.", "PASS"),
            ("15", ".sau olmayan dosyayı açma", "Arşiv dosyası uygunsuz veya bozuk mesajı.", "PASS"),
            ("16", "Olmayan .sau dosyasını açma", "Arşiv dosyası uygunsuz veya bozuk mesajı.", "PASS"),
            ("17", "-a için fazla parametre verme", "Arşiv dosyası uygunsuz veya bozuk mesajı.", "PASS"),
            ("18", "Sayısal olmayan başlıklı bozuk arşiv", "Bozuk arşiv olarak reddedilir.", "PASS"),
            ("19", "10 bayttan kısa arşiv", "Bozuk arşiv olarak reddedilir.", "PASS"),
            ("20", "Eksik metadata kaydı", "Bozuk arşiv olarak reddedilir.", "PASS"),
            ("21", "Payload boyutu metadata ile uyuşmayan arşiv", "Bozuk arşiv olarak reddedilir.", "PASS"),
            ("22", "Toplam 200 MB üzeri ASCII giriş", "200 MB sınırı mesajıyla reddedilir.", "PASS"),
        ],
        [0.38, 2.15, 3.1, 0.67],
    )
    add_para(doc, "Toplam sonuç: 22 senaryo çalıştırıldı, 22 başarılı, 0 başarısız.", bold=True, color=DARK_BLUE)

    add_heading(doc, "8. Dosya Yapısı")
    add_code_block(
        doc,
        """tarsau_odev/
  Makefile
  README.md
  tarsau.c
  testler/
    t1
    t2
    t3
    t4.txt
    t5.dat
  rapor/
    tarsau_proje_raporu.docx""",
    )

    add_heading(doc, "9. Sonuç")
    add_para(
        doc,
        "tarsau programı, proje föyünde tanımlanan sıkıştırmasız arşivleme davranışını C diliyle gerçekleştirmektedir. "
        "Başarılı akışlarda dosya içerikleri ve izinleri korunmuş, hata akışlarında program çökmeden anlamlı hata mesajlarıyla sonlanmıştır. "
        "Derleme, standart make test hedefi ve 22 uçtan uca test senaryosu başarıyla tamamlanmıştır."
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Ek A - Test Komut Özeti")
    add_para(doc, "Testler 19 Mayıs 2026 tarihinde yerel çalışma alanında, macOS üzerinde gcc ve make kullanılarak yürütülmüştür.")
    add_code_block(
        doc,
        """$ make clean && make
gcc -Wall -Wextra -std=c11 -O2 -o tarsau tarsau.c

$ make test
... diff kontrolleri hatasız tamamlandı

$ /tmp/tarsau_cases.../test-log.txt
SUMMARY pass=22 fail=0""",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
